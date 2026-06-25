"""
generate_docx.py — Genera un .docx con la plantilla del email en el escritorio.

Lee template.txt y crea Marsof-Email-Outreach.docx listo para copiar-pegar.
"""
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).parent
TEMPLATE = ROOT / "template.txt"
DESKTOP = Path(os.path.expanduser("~")) / "Desktop"
OUTPUT = DESKTOP / "Marsof - Email outreach.docx"


def load():
    raw = TEMPLATE.read_text(encoding="utf-8")
    subject = ""
    body_lines = []
    in_body = False
    for ln in raw.splitlines():
        if not in_body and ln.startswith("SUBJECT:"):
            subject = ln[len("SUBJECT:"):].strip()
            continue
        if ln.strip() == "---":
            in_body = True
            continue
        if in_body:
            body_lines.append(ln)
    return subject, "\n".join(body_lines).lstrip("\n")


def main():
    subject, body = load()

    doc = Document()

    # estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # título
    h = doc.add_heading("Email de outreach — Marsof / Saldea", level=1)

    # subtítulo informativo
    p = doc.add_paragraph()
    r = p.add_run("Copia este texto y pégalo en Zoho Mail al enviar a una gestoría / pyme. "
                  "Personaliza al menos 1 frase por email para evitar filtros de spam.")
    r.italic = True
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x7a)
    r.font.size = Pt(10)

    doc.add_paragraph()  # blank

    # asunto
    p = doc.add_paragraph()
    r = p.add_run("ASUNTO:")
    r.bold = True
    r.font.color.rgb = RGBColor(0x0e, 0xa5, 0xe9)
    p.add_run(f"  {subject}")

    doc.add_paragraph()  # blank

    # cuerpo
    p = doc.add_paragraph()
    r = p.add_run("CUERPO DEL EMAIL:")
    r.bold = True
    r.font.color.rgb = RGBColor(0x0e, 0xa5, 0xe9)

    for paragraph_text in body.split("\n\n"):
        p = doc.add_paragraph(paragraph_text.strip())
        p.paragraph_format.space_after = Pt(6)

    # firma
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Carlos Gálvez Carrillo").bold = True
    doc.add_paragraph("Fundador y CEO · Marsof Technology")

    # nota final
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Tip: cita algo concreto de la web del destinatario en la primera frase "
                  "(\"vi en vuestra web que también lleváis temas de Hacienda…\") para que "
                  "no parezca plantilla genérica.")
    r.italic = True
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x7a)
    r.font.size = Pt(10)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")


if __name__ == "__main__":
    main()
